"""Render overview.md → overview.docx with citation enrichment.

Best-effort converter: handles only the markdown shapes Scriptorium emits
(H1/H2/H3 headings, paragraphs, bullet/ordered lists, tables, inline
bold/italic/code, and `[paper_id:locator]` citations).

Failure must never block overview generation — the .md is the source of truth.
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[A-Za-z0-9_\-]+:[^\]]+\])"
)
_CITATION_RE = re.compile(r"\[([A-Za-z0-9_\-]+):([^\]]+)\]")


def render_overview_docx(md_path: Path, docx_path: Path, corpus_path: Path) -> None:
    """Render overview.md to .docx. Best-effort; caller isolates failures."""
    text = md_path.read_text(encoding="utf-8")
    body = _strip_frontmatter(text)
    corpus = _load_corpus(corpus_path)
    papers_dir = corpus_path.parent.parent / "sources" / "papers"
    ctx = {"corpus": corpus, "papers_dir": papers_dir, "misses": []}
    doc = Document()
    for block in _blocks(body):
        _render_block(doc, block, ctx)
    doc.save(str(docx_path))


def _strip_frontmatter(text: str) -> str:
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end != -1:
            return text[end + 5 :]
    return text


def _blocks(body: str) -> list[list[str]]:
    """Split body into blocks separated by blank lines."""
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in body.splitlines():
        if line.strip() == "":
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def _render_block(doc, block: list[str], ctx: dict) -> None:
    first = block[0]

    m = _HEADING_RE.match(first)
    if m and len(block) == 1:
        level = len(m.group(1))
        doc.add_heading(m.group(2).strip(), level=level)
        return

    if all(_BULLET_RE.match(line) for line in block):
        for line in block:
            p = doc.add_paragraph(style="List Bullet")
            _emit_runs(p, _BULLET_RE.match(line).group(1), ctx)
        return

    if all(_ORDERED_RE.match(line) for line in block):
        for line in block:
            p = doc.add_paragraph(style="List Number")
            _emit_runs(p, _ORDERED_RE.match(line).group(1), ctx)
        return

    if len(block) >= 2 and _TABLE_SEP_RE.match(block[1]) and "|" in block[0]:
        _render_table(doc, block)
        return

    para = doc.add_paragraph()
    _emit_runs(para, " ".join(line.strip() for line in block), ctx)


def _split_table_row(line: str) -> list[str]:
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def _render_table(doc, block: list[str]) -> None:
    header = _split_table_row(block[0])
    rows = [_split_table_row(line) for line in block[2:]]
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = header[i] if i < len(header) else ""
    for r, row in enumerate(rows, start=1):
        for i, cell in enumerate(table.rows[r].cells):
            cell.text = row[i] if i < len(row) else ""


def _emit_runs(paragraph, text: str, ctx: dict) -> None:
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif _CITATION_RE.fullmatch(part):
            _emit_citation(paragraph, part, ctx)
        else:
            paragraph.add_run(part)


def _load_corpus(corpus_path: Path) -> dict[str, dict]:
    if not corpus_path.exists() or corpus_path.stat().st_size == 0:
        return {}
    index: dict[str, dict] = {}
    for line in corpus_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            row = json.loads(line)
        except json.JSONDecodeError:
            continue
        pid = row.get("paper_id")
        if pid:
            index[pid] = row
    return index


def _first_author_surname(row: dict) -> str:
    authors = row.get("authors") or []
    if not authors:
        return row.get("paper_id", "")
    first = authors[0]
    if isinstance(first, str):
        return first.split(",")[0].strip()
    return first.get("family") or first.get("name", "")


_LOCATOR_SPACE_RE = re.compile(r"^(p\.|pp\.)(\S)")


def _normalize_locator(locator: str) -> str:
    """Insert a space between `p.`/`pp.` and digits: `p.12` → `p. 12`."""
    return _LOCATOR_SPACE_RE.sub(r"\1 \2", locator.strip())


def _citation_label(row: dict, locator: str) -> str:
    surname = _first_author_surname(row)
    year = row.get("year", "n.d.")
    return f"({surname} {year}, {_normalize_locator(locator)})"


def _citation_hyperlink(row: dict, papers_dir: Path) -> str | None:
    doi = row.get("doi")
    if doi:
        return f"https://doi.org/{doi}"
    url = row.get("url")
    if url:
        return url
    stub = papers_dir / f"{row.get('paper_id')}.md"
    if stub.exists():
        return str(stub)
    return None


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rPr.append(style)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _emit_citation(paragraph, raw: str, ctx: dict) -> None:
    m = _CITATION_RE.fullmatch(raw)
    pid, locator = m.group(1), m.group(2).strip()
    row = ctx["corpus"].get(pid)
    if not row:
        paragraph.add_run(raw)
        ctx["misses"].append(pid)
        return
    label = _citation_label(row, locator)
    link = _citation_hyperlink(row, ctx["papers_dir"])
    if link:
        _add_hyperlink(paragraph, link, label)
    else:
        paragraph.add_run(label)
