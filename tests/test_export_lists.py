"""Test: bullet and ordered lists render with correct docx styles."""
from pathlib import Path

from docx import Document

from scriptorium.export import render_overview_docx


def test_bullet_and_ordered_lists(tmp_path: Path):
    md = tmp_path / "overview.md"
    md.write_text(
        "- alpha\n- beta\n- gamma\n\n"
        "1. first\n2. second\n3. third\n",
        encoding="utf-8",
    )
    docx = tmp_path / "overview.docx"
    corpus = tmp_path / "corpus.jsonl"
    corpus.write_text("", encoding="utf-8")

    render_overview_docx(md, docx, corpus)

    doc = Document(str(docx))
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert styles.count("List Bullet") == 3
    assert styles.count("List Number") == 3
    texts = [p.text for p in doc.paragraphs]
    assert "alpha" in texts and "beta" in texts and "gamma" in texts
    assert "first" in texts and "second" in texts and "third" in texts
