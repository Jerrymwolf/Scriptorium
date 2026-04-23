"""Task 13: docx integrity smoke — valid zip + python-docx reopen."""
import zipfile
from pathlib import Path

from docx import Document

from scriptorium.export import render_overview_docx


def test_rendered_docx_is_valid_zip_and_reopens(tmp_path: Path):
    md = tmp_path / "overview.md"
    md.write_text(
        "# Title\n\nA paragraph with **bold** and [smith2024:p.1] citation.\n\n"
        "- one\n- two\n\n"
        "| A | B |\n| - | - |\n| 1 | 2 |\n",
        encoding="utf-8",
    )
    data = tmp_path / "data"
    data.mkdir()
    corpus = data / "corpus.jsonl"
    corpus.write_text(
        '{"paper_id":"smith2024","authors":["Smith, J."],"year":2024,"doi":"10/a"}\n',
        encoding="utf-8",
    )
    docx = tmp_path / "overview.docx"

    render_overview_docx(md, docx, corpus)

    assert zipfile.is_zipfile(docx)
    doc = Document(str(docx))
    assert any(p.style.name.startswith("Heading") for p in doc.paragraphs)
    assert len(doc.tables) == 1
