"""Test: fenced code and markdown-link URLs are exempt from citation enrichment."""
from pathlib import Path

from docx import Document

from scriptorium.export import render_overview_docx


def test_citation_inside_code_fence_left_alone(tmp_path: Path):
    md = tmp_path / "overview.md"
    md.write_text(
        "Regular [smith2024:p.1] here.\n\n"
        "```\n"
        "example syntax: [smith2024:p.1]\n"
        "```\n",
        encoding="utf-8",
    )
    corpus = tmp_path / "data" / "corpus.jsonl"
    corpus.parent.mkdir(parents=True, exist_ok=True)
    corpus.write_text(
        '{"paper_id":"smith2024","authors":["Smith, J."],"year":2024,"doi":"10.1/x"}\n',
        encoding="utf-8",
    )
    docx = tmp_path / "overview.docx"

    render_overview_docx(md, docx, corpus)

    doc = Document(str(docx))
    full = "\n".join(p.text for p in doc.paragraphs)
    # The code-fence line must be preserved verbatim (raw brackets).
    assert "example syntax: [smith2024:p.1]" in full
    # The non-fenced occurrence is enriched.
    assert "(Smith 2024, p. 1)" in full


def test_citation_inside_link_left_alone(tmp_path: Path):
    md = tmp_path / "overview.md"
    md.write_text(
        "See [link text](https://example.org/[smith2024:p.1]).\n",
        encoding="utf-8",
    )
    corpus = tmp_path / "data" / "corpus.jsonl"
    corpus.parent.mkdir(parents=True, exist_ok=True)
    corpus.write_text(
        '{"paper_id":"smith2024","authors":["Smith, J."],"year":2024,"doi":"10.1/x"}\n',
        encoding="utf-8",
    )
    docx = tmp_path / "overview.docx"

    render_overview_docx(md, docx, corpus)

    doc = Document(str(docx))
    # No (Smith 2024, p. 1) substitution inside the URL text.
    text = doc.paragraphs[0].text
    assert "[smith2024:p.1]" in text
