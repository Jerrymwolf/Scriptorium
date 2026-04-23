"""Test: `[paper_id:locator]` enriches to `(Author Year, locator)` with hyperlink.

The renderer derives `papers_dir` as `corpus_path.parent.parent / "sources" / "papers"`
so tests put corpus at `<tmp>/data/corpus.jsonl` and stubs at `<tmp>/sources/papers/…`
to match the real review layout.
"""
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from scriptorium.export import render_overview_docx


def _hyperlink_targets(paragraph) -> list[str]:
    """Return anchor URLs for every hyperlink in the paragraph."""
    urls = []
    part = paragraph.part
    for hl in paragraph._p.findall(qn("w:hyperlink")):
        rid = hl.get(qn("r:id"))
        if rid:
            urls.append(part.rels[rid].target_ref)
    return urls


def _write_corpus(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(json.dumps(r) for r in rows) + "\n", encoding="utf-8")


def test_citation_with_doi_becomes_doi_hyperlink(tmp_path: Path):
    md = tmp_path / "overview.md"
    md.write_text(
        "The claim holds [smith2024:p.12] under certain conditions.\n",
        encoding="utf-8",
    )
    corpus = tmp_path / "data" / "corpus.jsonl"
    _write_corpus(corpus, [{
        "paper_id": "smith2024",
        "authors": ["Smith, Jane"],
        "year": 2024,
        "doi": "10.1234/abc",
        "url": "https://example.org/smith2024",
    }])
    docx = tmp_path / "overview.docx"

    render_overview_docx(md, docx, corpus)

    doc = Document(str(docx))
    p = doc.paragraphs[0]
    assert "(Smith 2024, p. 12)" in p.text
    urls = _hyperlink_targets(p)
    assert urls == ["https://doi.org/10.1234/abc"]


def test_citation_falls_back_to_url_then_stub_then_plain(tmp_path: Path):
    corpus = tmp_path / "data" / "corpus.jsonl"
    md = tmp_path / "overview.md"
    docx = tmp_path / "overview.docx"

    # Row with url only → url used.
    _write_corpus(corpus, [{
        "paper_id": "jones2025",
        "authors": ["Jones, Lee"],
        "year": 2025,
        "url": "https://example.org/jones",
    }])
    md.write_text("See [jones2025:p.3].\n", encoding="utf-8")
    render_overview_docx(md, docx, corpus)
    doc = Document(str(docx))
    assert _hyperlink_targets(doc.paragraphs[0]) == ["https://example.org/jones"]

    # Row with no doi/url but stub exists → stub path used.
    _write_corpus(corpus, [{
        "paper_id": "lee2023",
        "authors": ["Lee, A."],
        "year": 2023,
    }])
    papers = tmp_path / "sources" / "papers"
    papers.mkdir(parents=True)
    stub = papers / "lee2023.md"
    stub.write_text("# lee2023\n")
    md.write_text("See [lee2023:p.1].\n", encoding="utf-8")
    render_overview_docx(md, docx, corpus)
    doc = Document(str(docx))
    assert _hyperlink_targets(doc.paragraphs[0]) == [str(stub)]

    # Row with no doi/url and no stub → plain text, no hyperlink.
    _write_corpus(corpus, [{
        "paper_id": "kim2022",
        "authors": ["Kim, B."],
        "year": 2022,
    }])
    md.write_text("See [kim2022:p.9].\n", encoding="utf-8")
    render_overview_docx(md, docx, corpus)
    doc = Document(str(docx))
    p = doc.paragraphs[0]
    assert "(Kim 2022, p. 9)" in p.text
    assert _hyperlink_targets(p) == []


def test_unknown_paper_id_left_raw(tmp_path: Path):
    corpus = tmp_path / "data" / "corpus.jsonl"
    _write_corpus(corpus, [])
    md = tmp_path / "overview.md"
    md.write_text("See [ghost2099:p.1].\n", encoding="utf-8")
    docx = tmp_path / "overview.docx"
    render_overview_docx(md, docx, corpus)
    doc = Document(str(docx))
    assert "[ghost2099:p.1]" in doc.paragraphs[0].text
