"""Test: markdown pipe tables render as docx tables."""
from pathlib import Path

from docx import Document

from scriptorium.export import render_overview_docx


def test_markdown_table_becomes_docx_table(tmp_path: Path):
    md = tmp_path / "overview.md"
    md.write_text(
        "| Paper | Claim |\n"
        "| --- | --- |\n"
        "| Smith 2024 | X improves Y |\n"
        "| Jones 2025 | X has no effect |\n",
        encoding="utf-8",
    )
    docx = tmp_path / "overview.docx"
    corpus = tmp_path / "corpus.jsonl"
    corpus.write_text("", encoding="utf-8")

    render_overview_docx(md, docx, corpus)

    doc = Document(str(docx))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.rows) == 3
    assert t.rows[0].cells[0].text == "Paper"
    assert t.rows[0].cells[1].text == "Claim"
    assert t.rows[1].cells[0].text == "Smith 2024"
    assert t.rows[2].cells[1].text == "X has no effect"
